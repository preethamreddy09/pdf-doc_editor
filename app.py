from flask import Flask, request, jsonify, send_file, render_template
import os
from werkzeug.utils import secure_filename
from docx import Document
import pdfplumber

app = Flask(__name__, static_folder="static", template_folder="templates")
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

# Serve the index.html file
@app.route("/")
def home():
    return render_template("index_edit.html")  # Load from the templates folder

# Extract text from uploaded file
@app.route("/extract_text", methods=["POST"])
def extract_text():
    file = request.files["file"]
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    file.save(filepath)

    text = ""
    if filename.endswith(".pdf"):
        with pdfplumber.open(filepath) as pdf:
            extracted_text = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    extracted_text.append(page_text)
            text = "\n".join(extracted_text)
    elif filename.endswith(".docx"):
        doc = Document(filepath)
        text = "\n".join([para.text for para in doc.paragraphs])
    else:
        return jsonify({"error": "Unsupported file format"}), 400

    return jsonify({"text": text, "filename": filename})


# Save edited text back to file
@app.route("/save_edited_text", methods=["POST"])
def save_edited_text():
    data = request.json
    edited_text = data["text"]
    filename = data["filename"]
    new_filename = filename.replace(".", "_edited.")

    new_filepath = os.path.join(app.config["UPLOAD_FOLDER"], new_filename)

    if filename.endswith(".pdf"):
        with open(new_filepath, "w", encoding="utf-8") as f:
            f.write(edited_text)  # Basic PDF saving, can improve with libraries
    elif filename.endswith(".docx"):
        doc = Document()
        doc.add_paragraph(edited_text)
        doc.save(new_filepath)
    else:
        return jsonify({"error": "Unsupported file format"}), 400

    return send_file(new_filepath, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)
